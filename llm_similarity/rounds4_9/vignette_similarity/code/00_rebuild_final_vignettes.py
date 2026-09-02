#!/usr/bin/env python3
"""Rebuild the final vignette DOCX with two minimal DG-C wording revisions."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "rating_runs" / "original" / "input" / "joint_action_vignettes_pre_dg_c_edit.docx"
OUTPUT = ROOT / "input" / "joint_action_vignettes_final.docx"
LOG = ROOT / "documentation" / "dg_c_wording_changes.json"

REPLACEMENTS = {
    "DG-C-P": {
        "old": (
            "Giovanni and Paola work together on a freelance project. Giovanni brings in the client, "
            "while Paola provides the technical expertise needed to finish the job. The client pays the "
            "entire completion bonus to Giovanni and leaves it to him to decide whether to share it. "
            "Giovanni is very happy about the payment. He thinks that the project has been a great win-win "
            "opportunity for him and Paola, and splits the profits from the joint collaboration with her, "
            "which still leaves him with a very good payoff."
        ),
        "new": (
            "Giovanni and Paola work together on a freelance project. Giovanni brings in the client, "
            "while Paola provides the technical expertise needed to finish the job. The client pays the "
            "entire completion bonus to Giovanni and leaves it to him to decide whether to share it. "
            "Giovanni views the project as a win-win collaboration made possible by their complementary "
            "contributions and shares the completion bonus with Paola."
        ),
    },
    "DG-C-K": {
        "old": (
            "Firm A and Firm B are matched through an online platform for a one-time production order. "
            "Firm A brings the customer, while Firm B supplies a component without which the order could "
            "not be completed. The platform pays the entire completion bonus to Firm A and leaves any "
            "transfer to B at A's discretion. The manager of Firm A is very pleased, the match has turned "
            "out to be a great value-creation opportunity for the two firms. She thus shares the bonus "
            "with Firm B, which leaves Firm A with a significant profit from the transaction."
        ),
        "new": (
            "Firm A and Firm B are matched through an online platform for a one-time production order. "
            "Firm A brings the customer, while Firm B supplies a component without which the order could "
            "not be completed. The platform pays the entire completion bonus to Firm A and leaves any "
            "transfer to B at A's discretion. The manager views the match as a value-creating collaboration "
            "made possible by the firms' complementary contributions and shares the completion bonus with "
            "Firm B."
        ),
    },
}


def main() -> None:
    document = Document(SOURCE)
    observed = set()
    for paragraph in document.paragraphs:
        for vignette_id, replacement in REPLACEMENTS.items():
            if paragraph.text == replacement["old"]:
                paragraph.text = replacement["new"]
                observed.add(vignette_id)
    if observed != set(REPLACEMENTS):
        raise ValueError(f"Expected both exact source paragraphs; replaced {sorted(observed)}")
    remaining = [
        vignette_id for vignette_id, replacement in REPLACEMENTS.items()
        if any(paragraph.text == replacement["old"] for paragraph in document.paragraphs)
    ]
    if remaining:
        raise AssertionError(f"Old wording remains for {remaining}")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    LOG.write_text(
        json.dumps(
            {
                "source": str(SOURCE.relative_to(ROOT)),
                "output": str(OUTPUT.relative_to(ROOT)),
                "purpose": (
                    "Remove incidental happiness and own-payoff language from both DG-C vignettes "
                    "while preserving their complementary-contribution and joint-value rationale."
                ),
                "changes": REPLACEMENTS,
            },
            ensure_ascii=False,
            indent=2,
        ) + "\n",
        encoding="utf-8",
    )
    print(f"Wrote {OUTPUT} with exactly two revised paragraphs")


if __name__ == "__main__":
    main()
